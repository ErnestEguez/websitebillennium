# -*- coding: utf-8 -*-
"""Genera Documentacion_ERP/Manual_Usuario_QuickInvoice.docx a partir de
manual_data.py (contenido) y las capturas en screenshots/quickinvoice/.
"""
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image

from manual_data import MODULOS

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(BASE_DIR)
SCREENSHOTS_DIR = os.path.join(ROOT_DIR, 'screenshots', 'quickinvoice')
OUTPUT_PATH = os.path.join(ROOT_DIR, 'Manual_Usuario_QuickInvoice.docx')

SECCIONES = [
    "Facturación y Ventas",
    "Inventario y Cartera por Cobrar",
    "Consultas, Compras y Cartera por Pagar",
    "Tesorería",
    "Contabilidad y Configuración",
    "Gerencia",
]

MAX_WIDTH_IN = 6.5
MAX_HEIGHT_IN = 8.0

MESES = [
    'enero', 'febrero', 'marzo', 'abril', 'mayo', 'junio',
    'julio', 'agosto', 'septiembre', 'octubre', 'noviembre', 'diciembre',
]


def add_toc(doc):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')

    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = r'TOC \o "1-3" \h \z \u'

    fld_sep = OxmlElement('w:fldChar')
    fld_sep.set(qn('w:fldCharType'), 'separate')

    placeholder = OxmlElement('w:t')
    placeholder.text = (
        'Haga clic derecho sobre esta área y seleccione '
        '"Actualizar campo" para generar la tabla de contenido.'
    )

    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')

    r = run._r
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)


def image_width_inches(path):
    with Image.open(path) as im:
        w_px, h_px = im.size
    aspect = h_px / w_px
    return min(MAX_WIDTH_IN, MAX_HEIGHT_IN / aspect)


def add_picture_centered(doc, path, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=Inches(image_width_inches(path)))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.font.size = Pt(9)
    cap_run.font.color.rgb = RGBColor(0x59, 0x59, 0x59)


def build():
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    # Portada
    title = doc.add_heading('Manual de Usuario', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('QuickInvoice ERP — Sistema de Gestión Empresarial')
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(16)

    company = doc.add_paragraph()
    company.alignment = WD_ALIGN_PARAGRAPH.CENTER
    company_run = company.add_run('Billennium Systems')
    company_run.font.size = Pt(12)

    from datetime import date
    hoy = date.today()
    fecha = doc.add_paragraph()
    fecha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fecha_run = fecha.add_run(f'{MESES[hoy.month - 1]} de {hoy.year}')
    fecha_run.font.size = Pt(11)
    fecha_run.italic = True

    doc.add_page_break()

    # Tabla de contenido
    doc.add_heading('Tabla de Contenido', level=1)
    add_toc(doc)
    doc.add_page_break()

    # Agrupar modulos por seccion preservando el orden de SECCIONES
    modulos_por_seccion = {s: [] for s in SECCIONES}
    for m in MODULOS:
        modulos_por_seccion[m['seccion']].append(m)
    for s in modulos_por_seccion:
        modulos_por_seccion[s].sort(key=lambda m: m['numero'])

    total_modulos = len(MODULOS)
    contador = 0

    for seccion in SECCIONES:
        doc.add_heading(seccion, level=1)
        for modulo in modulos_por_seccion[seccion]:
            contador += 1
            doc.add_heading(f"{modulo['numero']}. {modulo['titulo']}", level=2)

            desc_p = doc.add_paragraph(modulo['descripcion'])
            desc_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            img_path = os.path.join(SCREENSHOTS_DIR, modulo['archivo'] + '.png')
            if os.path.exists(img_path):
                caption = f"Figura {modulo['numero']}. {modulo['titulo']}"
                add_picture_centered(doc, img_path, caption)

            if modulo['elementos']:
                doc.add_heading('Elementos principales de la pantalla', level=3)
                for item in modulo['elementos']:
                    doc.add_paragraph(item, style='List Bullet')

            if modulo['pasos']:
                doc.add_heading('Cómo usarlo: paso a paso', level=3)
                for paso in modulo['pasos']:
                    doc.add_paragraph(paso, style='List Number')

            if modulo['notas']:
                doc.add_heading('Notas y advertencias', level=3)
                for nota in modulo['notas']:
                    doc.add_paragraph(nota, style='List Bullet')

            if contador < total_modulos:
                doc.add_page_break()

    doc.save(OUTPUT_PATH)
    print(f'Generado: {OUTPUT_PATH}')


if __name__ == '__main__':
    build()
