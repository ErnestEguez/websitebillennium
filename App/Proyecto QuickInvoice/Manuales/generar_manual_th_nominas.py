"""
Genera el Manual de Usuario: Talento Humano y Nóminas – QuickInvoice
Ejecutar: python generar_manual_th_nominas.py
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, datetime

# ─── Helpers ────────────────────────────────────────────────────────────────

def set_col_widths(table, widths_cm):
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = Cm(widths_cm[j])

def shade_cell(cell, fill_hex):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex)
    tcPr.append(shd)

def add_table_border(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ['top','left','bottom','right','insideH','insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),   'single')
        el.set(qn('w:sz'),    '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'AAAAAA')
        borders.append(el)
    tblPr.append(borders)

AZUL_OSCURO  = '1E3A5F'
AZUL_MEDIO   = '2E5F9E'
AZUL_CLARO   = 'D6E4F7'
VERDE        = '1B5E20'
VERDE_CLARO  = 'E8F5E9'
GRIS_CABEZA  = '37474F'
GRIS_CLARO   = 'F5F5F5'
BLANCO       = 'FFFFFF'
AMARILLO     = 'FFF9C4'
NARANJA_TXT  = 'E65100'

# ─── Document setup ─────────────────────────────────────────────────────────

doc = Document()

# Page margins
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

# ─── Styles ─────────────────────────────────────────────────────────────────

def style_h1(para, text):
    para.clear()
    run = para.add_run(text)
    run.font.size  = Pt(18)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    para.paragraph_format.space_before = Pt(24)
    para.paragraph_format.space_after  = Pt(6)

def style_h2(para, text):
    para.clear()
    run = para.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x2E, 0x5F, 0x9E)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)

def style_h3(para, text):
    para.clear()
    run = para.add_run(text)
    run.font.size  = Pt(12)
    run.font.bold  = True
    run.font.color.rgb = RGBColor(0x37, 0x47, 0x4F)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after  = Pt(3)

def add_h1(doc, text):
    p = doc.add_paragraph()
    style_h1(p, text)
    return p

def add_h2(doc, text):
    p = doc.add_paragraph()
    style_h2(p, text)
    return p

def add_h3(doc, text):
    p = doc.add_paragraph()
    style_h3(p, text)
    return p

def add_body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Cm(1 + level * 0.5)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p

def add_nota(doc, text, color=AMARILLO, label='NOTA'):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    shade_cell(cell, color)
    cell.paragraphs[0].clear()
    r1 = cell.paragraphs[0].add_run(f'{label}: ')
    r1.font.bold = True
    r1.font.size = Pt(10)
    r2 = cell.paragraphs[0].add_run(text)
    r2.font.size = Pt(10)
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    add_table_border(tbl)
    doc.add_paragraph()

def add_image_placeholder(doc, label, ancho='16 cm', alto='8 cm'):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    shade_cell(cell, 'EEF2FF')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'[ IMAGEN: {label} ]\nTamaño sugerido: {ancho} × {alto}')
    r.font.size  = Pt(9)
    r.font.color.rgb = RGBColor(0x5C, 0x6B, 0xC0)
    r.font.italic = True
    add_table_border(tbl)
    doc.add_paragraph()

def header_table_row(table, row_idx, texts, fill=AZUL_OSCURO):
    row = table.rows[row_idx]
    for i, txt in enumerate(texts):
        if i < len(row.cells):
            cell = row.cells[i]
            shade_cell(cell, fill)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            run.font.bold  = True
            run.font.size  = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

def data_row(table, row_idx, texts, fill=None):
    row = table.rows[row_idx]
    bg  = fill or (GRIS_CLARO if row_idx % 2 == 0 else BLANCO)
    for i, txt in enumerate(texts):
        if i < len(row.cells):
            cell = row.cells[i]
            shade_cell(cell, bg)
            cell.paragraphs[0].clear()
            run = cell.paragraphs[0].add_run(txt)
            run.font.size = Pt(10)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

def page_break(doc):
    doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  PORTADA
# ════════════════════════════════════════════════════════════════════════════

add_image_placeholder(doc, 'Logo de Billennium / QuickInvoice (colocar aquí)', '10 cm', '3 cm')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('MANUAL DE USUARIO')
r.font.size  = Pt(28)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Módulo de Talento Humano y Nóminas')
r.font.size  = Pt(20)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x2E, 0x5F, 0x9E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('QuickInvoice — Sistema de Gestión Empresarial')
r.font.size  = Pt(13)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Versión 1.0  ·  {datetime.date.today().strftime("%B %Y")}')
r.font.size  = Pt(11)
r.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Billennium Systems')
r.font.size  = Pt(11)
r.font.bold  = True

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  TABLA DE CONTENIDOS (manual)
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, 'TABLA DE CONTENIDOS')

toc = [
    ('1.', 'Introducción y Alcance del Módulo', '4'),
    ('2.', 'Configuración Inicial (antes de operar)', '5'),
    ('  2.1', 'Parámetros de Nómina', '5'),
    ('  2.2', 'Conceptos de Nómina', '6'),
    ('  2.3', 'Cuentas Contables de Nómina', '8'),
    ('  2.4', 'Estructura Organizativa', '9'),
    ('3.', 'Módulo de Talento Humano', '10'),
    ('  3.1', 'Dashboard de Talento', '10'),
    ('  3.2', 'Reclutamiento y Selección (Vacantes)', '11'),
    ('  3.3', 'Gestión de Empleados', '14'),
    ('  3.4', 'Onboarding y Offboarding (Checklists)', '17'),
    ('  3.5', 'Evaluación de Desempeño', '19'),
    ('  3.6', 'Capacitación', '21'),
    ('  3.7', 'Clima Organizacional', '23'),
    ('  3.8', 'Finiquitos', '25'),
    ('4.', 'Módulo de Liquidación de Nóminas', '27'),
    ('  4.1', 'Períodos de Nómina', '27'),
    ('  4.2', 'Rol de Pagos', '29'),
    ('  4.3', 'Anticipos de Quincena', '34'),
    ('  4.4', 'Novedades y Descuentos Recurrentes', '36'),
    ('  4.5', 'Liquidación de Décimos (13.º y 14.º)', '39'),
    ('  4.6', 'Liquidación de Vacaciones', '41'),
    ('  4.7', 'Estado Económico / Capacidad de Pago', '43'),
    ('  4.8', 'Reportes de Nómina', '44'),
    ('5.', 'Contabilidad Automática de Nómina', '45'),
    ('6.', 'Permisos de Usuario', '48'),
    ('7.', 'Flujo de Trabajo Mensual Recomendado', '50'),
    ('8.', 'Preguntas Frecuentes', '51'),
]

tbl = doc.add_table(rows=len(toc), cols=3)
set_col_widths(tbl, [1.5, 11.5, 1.5])
for i, (num, titulo, pag) in enumerate(toc):
    bg = AZUL_CLARO if not num.startswith(' ') else BLANCO
    shade_cell(tbl.rows[i].cells[0], bg)
    shade_cell(tbl.rows[i].cells[1], bg)
    shade_cell(tbl.rows[i].cells[2], bg)
    for j, txt in enumerate([num, titulo, pag]):
        cell = tbl.rows[i].cells[j]
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(10)
        run.font.bold = not num.startswith(' ')
        cell.paragraphs[0].paragraph_format.space_after = Pt(1)
    tbl.rows[i].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
add_table_border(tbl)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  1. INTRODUCCIÓN
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, '1. Introducción y Alcance del Módulo')

add_body(doc,
    'El módulo de Talento Humano y Nóminas de QuickInvoice centraliza toda la gestión del '
    'capital humano de la empresa: desde el reclutamiento y selección de personal, pasando '
    'por la administración de empleados, onboarding, evaluaciones de desempeño, capacitación '
    'y clima laboral, hasta el cálculo y contabilización automática del rol de pagos, décimos, '
    'vacaciones, anticipos y finiquitos.')

add_body(doc,
    'Está diseñado para cumplir con la legislación laboral ecuatoriana, incluyendo el Código '
    'del Trabajo, las disposiciones del IESS y la normativa del SRI, automatizando los '
    'cálculos de IESS patronal y personal, fondos de reserva, decimotercero, decimocuarto, '
    'vacaciones y retenciones de impuesto a la renta.')

add_h2(doc, 'Componentes del módulo')

componentes = [
    ('Talento Humano', 'Estructura organizativa · Vacantes/Reclutamiento · Empleados · '
                       'Onboarding/Offboarding · Evaluación de Desempeño · Capacitación · '
                       'Clima Organizacional · Finiquitos · Dashboard'),
    ('Liquidación de Nóminas', 'Períodos · Rol de Pagos · Anticipos · Novedades/Préstamos · '
                                'Décimos 13.º y 14.º · Vacaciones · Reportes · Estado Económico'),
    ('Contabilidad Automática', 'Asientos contables de rol mensual · Provisión de leyes sociales · '
                                 'Anticipos · Registro en el sistema LP/LedgerPro integrado'),
]

tbl = doc.add_table(rows=len(componentes)+1, cols=2)
set_col_widths(tbl, [5, 10.5])
header_table_row(tbl, 0, ['Componente', 'Incluye'])
for i, (comp, det) in enumerate(componentes):
    data_row(tbl, i+1, [comp, det])
add_table_border(tbl)

doc.add_paragraph()
add_nota(doc,
    'Este módulo requiere que la empresa tenga configurado el Catálogo de Cuentas Contables '
    'en LedgerPro para que la contabilización automática del rol funcione correctamente.')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  2. CONFIGURACIÓN INICIAL
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, '2. Configuración Inicial')

add_body(doc,
    'Antes de operar el módulo es necesario completar los siguientes pasos en el orden '
    'indicado. Una configuración incorrecta puede afectar todos los cálculos de nómina.')

# 2.1 Parámetros
add_h2(doc, '2.1 Parámetros de Nómina')

add_body(doc,
    'Ruta: Liquidación de Nóminas → Parámetros de Nómina')
add_body(doc,
    'Los parámetros definen los porcentajes legales vigentes que el sistema usa para calcular '
    'automáticamente los aportes al IESS, fondo de reserva y otras obligaciones patronales.')

add_image_placeholder(doc, 'Pantalla: Parámetros de Nómina — vista general')

params = [
    ('Salario Básico Unificado (SBU)', 'Valor vigente según resolución anual del MRL. Base para cálculo del 14.º.', 'USD 460.00 (2024)'),
    ('Aporte Personal IESS', 'Porcentaje descontado al empleado para el IESS.', '9.45 %'),
    ('Aporte Patronal IESS', 'Porcentaje que aporta el empleador al IESS.', '11.15 %'),
    ('Fondo de Reserva', 'Porcentaje acumulado a partir del 2.º año de trabajo.', '8.33 %'),
    ('Décimo Tercero', 'Porcentaje de provisión mensual sobre remuneración.', '8.33 %'),
    ('Décimo Cuarto', 'Porcentaje de provisión mensual sobre SBU/12.', '8.33 % del SBU'),
    ('Vacaciones', 'Porcentaje de provisión mensual sobre remuneración.', '4.17 %'),
]

tbl = doc.add_table(rows=len(params)+1, cols=3)
set_col_widths(tbl, [5, 7, 3.5])
header_table_row(tbl, 0, ['Parámetro', 'Descripción', 'Valor referencial'])
for i, row in enumerate(params):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()
add_nota(doc,
    'Actualizar el SBU cada enero según el Acuerdo Ministerial emitido por el Ministerio del '
    'Trabajo. El cambio afecta inmediatamente a todos los roles que se generen desde ese momento.', AMARILLO)

# 2.2 Conceptos
add_h2(doc, '2.2 Conceptos de Nómina')

add_body(doc,
    'Ruta: Liquidación de Nóminas → Conceptos de Nómina')
add_body(doc,
    'Los conceptos son los rubros que componen el rol: ingresos (sueldo, horas extras, '
    'comisiones…) y descuentos (IESS personal, préstamos, multas…). El sistema incluye '
    'los conceptos estándar ecuatorianos precargados.')

add_image_placeholder(doc, 'Pantalla: Conceptos de Nómina — listado')

add_h3(doc, 'Conceptos estándar precargados')

conceptos = [
    ('SUELDO',        'ingreso',    'Sueldo mensual base',                              'Sí', 'Sí', 'Siempre'),
    ('HRS_EXTRA_50',  'ingreso',    'Horas suplementarias (recargo 50%)',                'Sí', 'Sí', 'No'),
    ('HRS_EXTRA_100', 'ingreso',    'Horas extraordinarias (recargo 100%)',              'Sí', 'Sí', 'No'),
    ('HRS_NOCT_25',   'ingreso',    'Horas nocturnas (recargo 25%)',                     'Sí', 'Sí', 'No'),
    ('COMISION',      'ingreso',    'Comisiones variables',                              'Sí', 'Sí', 'No'),
    ('IESS_PERSONAL', 'descuento',  'Aporte personal IESS 9.45%',                       'Sí', 'No', 'Siempre'),
    ('ANTICIPO',      'descuento',  'Descuento de anticipo de quincena',                 'No', 'No', 'No'),
    ('PRESTAMO',      'descuento',  'Cuota de préstamo/novedad',                        'No', 'No', 'No'),
    ('MULTA',         'descuento',  'Multas o sanciones',                               'No', 'No', 'No'),
    ('RETENCION_RF',  'descuento',  'Retención en la fuente de impuesto a la renta',    'No', 'No', 'No'),
]

tbl = doc.add_table(rows=len(conceptos)+1, cols=6)
set_col_widths(tbl, [2.5, 2.0, 5.5, 1.3, 1.3, 1.8])
header_table_row(tbl, 0, ['Código', 'Tipo', 'Descripción', 'Afecta IESS', 'Afecta Renta', 'Aplica siempre'])
for i, row in enumerate(conceptos):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()
add_h3(doc, 'Fórmulas de horas extras/nocturnas (Ecuador)')

add_body(doc, 'El sistema calcula automáticamente el monto al ingresar las horas trabajadas:')

formulas = [
    ('HRS_EXTRA_50',  '(Sueldo ÷ 240) × 1.5 × horas',   'Horas suplementarias — incluye sueldo base + 50% de recargo. Divisor 240 = 8h/día × 30 días (Art. 47 CT).'),
    ('HRS_EXTRA_100', '(Sueldo ÷ 240) × 2.0 × horas',   'Horas extraordinarias — incluye sueldo base + 100% de recargo.'),
    ('HRS_NOCT_25',   '(Sueldo ÷ 240) × 0.25 × horas',  'Solo el recargo nocturno del 25%. El sueldo base ya está incluido en el concepto SUELDO.'),
]

tbl = doc.add_table(rows=len(formulas)+1, cols=3)
set_col_widths(tbl, [2.5, 5, 8])
header_table_row(tbl, 0, ['Concepto', 'Fórmula', 'Nota'])
for i, row in enumerate(formulas):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()
add_nota(doc,
    'Para agregar el concepto HRS_NOCT_25 si el catálogo ya estaba creado, use el botón '
    '"Nuevo Concepto" en la pantalla de Conceptos de Nómina e ingrese los valores indicados.', AMARILLO)

# 2.3 Cuentas contables
add_h2(doc, '2.3 Cuentas Contables de Nómina')

add_body(doc, 'Ruta: Liquidación de Nóminas → Cuentas Contables Nómina')
add_body(doc,
    'Mapea cada rubro de nómina con su cuenta contable correspondiente en el plan de cuentas '
    'de LedgerPro. Sin este mapeo la contabilización automática no puede ejecutarse.')

add_image_placeholder(doc, 'Pantalla: Cuentas Contables de Nómina')

cuentas = [
    ('Gasto Sueldos',          'Cuenta de gastos donde se carga el sueldo bruto del personal.',               'Ej.: 5.1.01.001'),
    ('Gasto IESS Patronal',    'Gasto por el aporte patronal al IESS.',                                        'Ej.: 5.1.02.001'),
    ('Gasto Décimo Tercero',   'Provisión mensual del decimotercero.',                                          'Ej.: 5.1.03.001'),
    ('Gasto Décimo Cuarto',    'Provisión mensual del decimocuarto.',                                           'Ej.: 5.1.03.002'),
    ('Gasto Vacaciones',       'Provisión mensual de vacaciones.',                                              'Ej.: 5.1.03.003'),
    ('Gasto Fondo de Reserva', 'Provisión mensual de fondo de reserva.',                                       'Ej.: 5.1.03.004'),
    ('IESS por Pagar',         'Pasivo — total IESS (personal + patronal) pendiente de pago.',                 'Ej.: 2.1.03.001'),
    ('Sueldos por Pagar',      'Pasivo — neto a pagar a empleados.',                                           'Ej.: 2.1.02.001'),
    ('Prov. Décimo Tercero',   'Pasivo acumulado — provisión 13.º sueldo.',                                    'Ej.: 2.1.04.001'),
    ('Prov. Décimo Cuarto',    'Pasivo acumulado — provisión 14.º sueldo.',                                    'Ej.: 2.1.04.002'),
    ('Prov. Vacaciones',       'Pasivo acumulado — provisión vacaciones.',                                     'Ej.: 2.1.04.003'),
    ('Prov. Fondo de Reserva', 'Pasivo acumulado — provisión fondo de reserva.',                               'Ej.: 2.1.04.004'),
    ('Anticipo Empleados',     'Activo — anticipos pagados pendientes de descuento.',                          'Ej.: 1.1.05.001'),
]

tbl = doc.add_table(rows=len(cuentas)+1, cols=3)
set_col_widths(tbl, [4.5, 8, 3])
header_table_row(tbl, 0, ['Campo', 'Descripción', 'Ejemplo'])
for i, row in enumerate(cuentas):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

# 2.4 Estructura organizativa
add_h2(doc, '2.4 Estructura Organizativa')

add_body(doc, 'Ruta: Talento Humano → Estructura Organizativa')
add_body(doc,
    'Define los departamentos (secciones) y los cargos disponibles en la empresa. '
    'Esta información se usa en las fichas de empleados, evaluaciones, checklists y reportes.')

add_image_placeholder(doc, 'Pantalla: Estructura Organizativa — secciones y cargos')

add_bullet(doc, 'Sección: departamento o área de la empresa (Ej.: Contabilidad, Ventas, TI).')
add_bullet(doc, 'Cargo: puesto dentro de una sección (Ej.: Contador General, Asesor Comercial).')
add_bullet(doc, 'Un cargo pertenece a una sección y puede tener un cargo superior (jerarquía).')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  3. MÓDULO DE TALENTO HUMANO
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, '3. Módulo de Talento Humano')

# 3.1 Dashboard
add_h2(doc, '3.1 Dashboard de Talento')

add_body(doc, 'Ruta: Talento Humano → Dashboard')
add_body(doc,
    'Vista ejecutiva con los indicadores clave del área de Recursos Humanos. Permite tomar '
    'decisiones rápidas sin necesidad de ingresar a cada módulo específico.')

add_image_placeholder(doc, 'Pantalla: Dashboard de Talento — indicadores y gráficos')

indicadores = [
    ('Total de empleados activos', 'Número de empleados vigentes en la empresa.'),
    ('Vacantes abiertas',          'Cantidad de posiciones en proceso de reclutamiento.'),
    ('Anticipos del período',      'Monto total de anticipos emitidos en el período activo.'),
    ('Onboarding pendiente',       'Checklists de bienvenida en proceso de completar.'),
    ('Evaluaciones programadas',   'Evaluaciones de desempeño pendientes de respuesta.'),
    ('Capacitaciones en curso',    'Cursos activos con empleados inscritos.'),
    ('Puntuación de clima',        'Promedio de satisfacción de la última encuesta.'),
    ('Finiquitos del mes',         'Liquidaciones finales procesadas en el mes actual.'),
]

tbl = doc.add_table(rows=len(indicadores)+1, cols=2)
set_col_widths(tbl, [6, 9.5])
header_table_row(tbl, 0, ['Indicador', 'Descripción'])
for i, row in enumerate(indicadores):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

# 3.2 Reclutamiento
add_h2(doc, '3.2 Reclutamiento y Selección')

add_body(doc, 'Ruta: Talento Humano → Reclutamiento (Vacantes)')
add_body(doc,
    'Gestión completa del proceso de reclutamiento desde la publicación de una vacante hasta '
    'la contratación del candidato. Incluye evaluación automática de CVs con Inteligencia Artificial.')

add_image_placeholder(doc, 'Pantalla: Lista de Vacantes — estados y filtros')

add_h3(doc, 'Crear una Vacante')

add_body(doc, 'Haga clic en "Nueva Vacante" y complete:')

vacante_campos = [
    ('Título del puesto',     'Nombre visible del cargo que se busca cubrir.', 'Obligatorio'),
    ('Sección / Cargo',       'Área y posición dentro de la estructura organizativa.', 'Obligatorio'),
    ('Modalidad de trabajo',  'Presencial / Remoto / Híbrido.', 'Obligatorio'),
    ('Descripción del puesto','Responsabilidades y funciones principales.', 'Recomendado'),
    ('Requisitos mínimos',    'Educación, experiencia y competencias requeridas.', 'Recomendado'),
    ('Rango salarial',        'Rango de remuneración mensual ofrecido.', 'Opcional'),
    ('Fecha límite',          'Fecha hasta la que se reciben postulaciones.', 'Opcional'),
    ('Estado',                'Abierta / En proceso / Cerrada / Cancelada.', 'Automático'),
]

tbl = doc.add_table(rows=len(vacante_campos)+1, cols=3)
set_col_widths(tbl, [4, 8, 3])
header_table_row(tbl, 0, ['Campo', 'Descripción', 'Requerido'])
for i, row in enumerate(vacante_campos):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

add_h3(doc, 'Gestión de Candidatos')

add_body(doc,
    'Dentro de cada vacante se registran los candidatos y se hace seguimiento de su proceso:')

etapas = [
    ('Postulado',    'El candidato ha aplicado. Se adjunta el CV (PDF/DOC).'),
    ('Preseleccionado', 'El candidato pasó el filtro inicial de requisitos.'),
    ('Entrevista',   'Se programó/realizó la entrevista.'),
    ('Evaluación',   'El candidato está en proceso de pruebas técnicas o psicológicas.'),
    ('Oferta',       'Se le presentó una oferta económica al candidato.'),
    ('Contratado',   'Candidato aceptó y se inicia el proceso de vinculación.'),
    ('Descartado',   'El candidato no continuó en el proceso.'),
]

tbl = doc.add_table(rows=len(etapas)+1, cols=2)
set_col_widths(tbl, [4, 11.5])
header_table_row(tbl, 0, ['Etapa', 'Descripción'])
for i, row in enumerate(etapas):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

add_h3(doc, 'Evaluación de Candidatos con IA')

add_body(doc,
    'El sistema utiliza Inteligencia Artificial (Gemini) para analizar el CV del candidato '
    'y compararlo con los requisitos de la vacante, generando un puntaje y comentarios.')

add_image_placeholder(doc, 'Pantalla: Evaluación IA — puntaje, fortalezas, debilidades y preguntas sugeridas')

add_body(doc, 'El análisis IA devuelve:')
add_bullet(doc, 'Puntuación de 0 a 100 indicando el porcentaje de ajuste al perfil solicitado.')
add_bullet(doc, 'Fortalezas del candidato en relación con los requisitos del puesto.')
add_bullet(doc, 'Brechas o áreas de mejora que deberían abordarse en la entrevista.')
add_bullet(doc, '5 preguntas de entrevista personalizadas al perfil del candidato.')
add_bullet(doc, 'Recomendación final: Altamente recomendado / Recomendado / No recomendado.')

add_nota(doc,
    'Para que la IA analice al candidato: 1) Adjunte el CV en la sección de documentos del candidato, '
    '2) Haga clic en el botón "Evaluar con IA", 3) Espere 10-30 segundos mientras Gemini procesa el documento. '
    'El análisis se guarda automáticamente y puede re-ejecutarse en cualquier momento.', VERDE_CLARO)

# 3.3 Empleados
add_h2(doc, '3.3 Gestión de Empleados')

add_body(doc, 'Ruta: Talento Humano → Empleados')
add_body(doc,
    'Administra el ciclo de vida completo del empleado desde su ingreso hasta su salida. '
    'Incluye datos personales, laborales, bancarios, foto, historial salarial y configuración de anticipos.')

add_image_placeholder(doc, 'Pantalla: Lista de Empleados — búsqueda, filtros y acciones')

add_h3(doc, 'Datos de la Ficha del Empleado')

fichas = {
    'Pestaña Personal': [
        ('Nombres y Apellidos', 'Nombre completo del empleado.', 'Obligatorio'),
        ('Cédula / Pasaporte',  'Número de identificación único.',             'Obligatorio'),
        ('Fecha de nacimiento', 'Para cálculos de edad y beneficios.',         'Obligatorio'),
        ('Género',              'Masculino / Femenino / Otro.',                'Obligatorio'),
        ('Estado civil',        'Soltero, casado, divorciado, etc.',           'Opcional'),
        ('Dirección',           'Domicilio actual del empleado.',              'Opcional'),
        ('Teléfonos / Email',   'Datos de contacto directo.',                  'Opcional'),
        ('Foto',                'Fotografía del empleado (JPEG/PNG).',         'Opcional'),
    ],
    'Pestaña Laboral': [
        ('Cargo',               'Puesto dentro de la estructura organizativa.', 'Obligatorio'),
        ('Sección',             'Departamento al que pertenece.',              'Obligatorio'),
        ('Fecha de ingreso',    'Inicio de la relación laboral.',              'Obligatorio'),
        ('Fecha de salida',     'Se llena al momento del finiquito.',          'Automático'),
        ('Tipo de contrato',    'Indefinido, fijo, a prueba, por obra, etc.', 'Obligatorio'),
        ('Sueldo base',         'Remuneración mensual vigente.',               'Obligatorio'),
        ('Estado',              'Activo / Inactivo.',                          'Automático'),
    ],
    'Pestaña Bancaria': [
        ('Banco',               'Entidad bancaria donde recibe el pago.',     'Obligatorio'),
        ('Tipo de cuenta',      'Ahorros / Corriente.',                       'Obligatorio'),
        ('Número de cuenta',    'Número completo de la cuenta bancaria.',     'Obligatorio'),
    ],
    'Pestaña Anticipo': [
        ('Tipo de anticipo',    'Porcentaje del sueldo o monto fijo.',        'Opcional'),
        ('Valor predeterminado','% o monto usado al generar el anticipo.',    'Opcional'),
    ],
    'Pestaña Emergencia': [
        ('Contacto emergencia', 'Nombre y relación del contacto.',            'Recomendado'),
        ('Teléfono emergencia', 'Número directo del contacto.',               'Recomendado'),
    ],
}

for pestaña, campos in fichas.items():
    add_h3(doc, pestaña)
    tbl = doc.add_table(rows=len(campos)+1, cols=3)
    set_col_widths(tbl, [4, 8, 3])
    header_table_row(tbl, 0, ['Campo', 'Descripción', 'Requerido'])
    for i, row in enumerate(campos):
        data_row(tbl, i+1, list(row))
    add_table_border(tbl)
    doc.add_paragraph()

add_image_placeholder(doc, 'Pantalla: Ficha del Empleado — pestañas Personal / Laboral / Bancario')

add_h3(doc, 'Historial de Salarios')

add_body(doc,
    'Cada vez que se modifica el sueldo base de un empleado, el sistema registra '
    'automáticamente el historial con fecha de cambio, valor anterior, valor nuevo y usuario que realizó el cambio. '
    'Visible en la pestaña correspondiente de la ficha del empleado.')

add_image_placeholder(doc, 'Pantalla: Historial de Salarios del Empleado')

# 3.4 Onboarding
add_h2(doc, '3.4 Onboarding y Offboarding')

add_body(doc, 'Rutas: Talento Humano → Plantillas Checklist  |  Talento Humano → Onboarding/Offboarding')
add_body(doc,
    'Sistema de checklists para garantizar que todos los procesos de ingreso y salida de '
    'personal se completen correctamente, sin olvidar ningún paso crítico.')

add_image_placeholder(doc, 'Pantalla: Plantillas de Checklist — lista y modal de creación')

add_h3(doc, 'Plantillas de Checklist')

add_body(doc, 'Primero se crean las plantillas que definirán las tareas a ejecutar:')
add_bullet(doc, 'Tipo: Onboarding (bienvenida) o Offboarding (salida).')
add_bullet(doc, 'Aplicable a: Todos los empleados, un cargo específico, o una sección.')
add_bullet(doc, 'Tareas: listado de pasos con descripción, área responsable y días de plazo.')

add_body(doc, 'Ejemplos de tareas de Onboarding:')
add_bullet(doc, 'Entrega de equipo de trabajo (laptop, celular) — Responsable: TI — Plazo: día 1', 1)
add_bullet(doc, 'Creación de usuario en sistemas — Responsable: TI — Plazo: día 1', 1)
add_bullet(doc, 'Inducción a la empresa y políticas — Responsable: RRHH — Plazo: día 1-3', 1)
add_bullet(doc, 'Firma de contrato y documentos legales — Responsable: Legal — Plazo: día 1', 1)
add_bullet(doc, 'Afiliación al IESS — Responsable: Nómina — Plazo: antes del primer día', 1)

add_h3(doc, 'Ejecutar un Checklist')

add_body(doc,
    'Al vincular o desvincular un empleado, se genera automáticamente un checklist basado '
    'en la plantilla aplicable. El responsable de cada área puede marcar las tareas como completadas.')

add_image_placeholder(doc, 'Pantalla: Checklist de Onboarding — tareas y progreso')

add_bullet(doc, 'Estado: Pendiente → En progreso → Completado.')
add_bullet(doc, 'Cada tarea se puede asignar a un responsable diferente.')
add_bullet(doc, 'El sistema registra quién completó cada tarea y cuándo.')
add_bullet(doc, 'El porcentaje de avance se muestra en el Dashboard de Talento.')

# 3.5 Evaluación desempeño
add_h2(doc, '3.5 Evaluación de Desempeño')

add_body(doc, 'Ruta: Talento Humano → Evaluación de Desempeño')
add_body(doc,
    'Permite ejecutar ciclos de evaluación del desempeño de los empleados mediante '
    'diferentes modalidades y con criterios configurables por la empresa.')

add_image_placeholder(doc, 'Pantalla: Períodos de Evaluación y lista de evaluaciones')

modalidades = [
    ('Por supervisor',   'El jefe directo evalúa al empleado según los criterios definidos.'),
    ('Autoevaluación',   'El empleado se evalúa a sí mismo.'),
    ('Evaluación 360°',  'El empleado es evaluado por su jefe, pares y subordinados de forma anónima.'),
]

tbl = doc.add_table(rows=len(modalidades)+1, cols=2)
set_col_widths(tbl, [4, 11.5])
header_table_row(tbl, 0, ['Modalidad', 'Descripción'])
for i, row in enumerate(modalidades):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()

add_h3(doc, 'Proceso de Evaluación')

pasos = [
    'Crear un Período de Evaluación: defina nombre, fechas de apertura y cierre, y tipo de evaluación.',
    'Configurar criterios de evaluación: lista de competencias o KPIs a evaluar.',
    'Generar evaluaciones: el sistema crea automáticamente los formularios para cada empleado/evaluador.',
    'Los evaluadores completan el formulario en línea con puntuaciones y comentarios.',
    'Al cierre del período, se generan los informes consolidados por empleado y por departamento.',
    'Los resultados quedan en el historial del empleado para seguimiento y planes de desarrollo.',
]

for i, paso in enumerate(pasos):
    add_bullet(doc, f'Paso {i+1}: {paso}')

add_image_placeholder(doc, 'Pantalla: Formulario de Evaluación — criterios y puntajes')

# 3.6 Capacitación
add_h2(doc, '3.6 Capacitación')

add_body(doc, 'Ruta: Talento Humano → Capacitación')
add_body(doc,
    'Registro y control de cursos de capacitación, tanto internos como externos, '
    'con seguimiento de las inscripciones y certificados de cada empleado.')

add_image_placeholder(doc, 'Pantalla: Capacitaciones — lista de cursos e inscripciones')

campos_curso = [
    ('Nombre del curso',  'Título completo de la capacitación.', 'Obligatorio'),
    ('Tipo',              'Interno (dictado por la empresa) o Externo (proveedor externo).', 'Obligatorio'),
    ('Modalidad',         'Presencial / Virtual / Mixto.', 'Obligatorio'),
    ('Instructor',        'Nombre del instructor o proveedor.', 'Opcional'),
    ('Fechas',            'Fecha de inicio y fin del curso.', 'Obligatorio'),
    ('Horas',             'Duración total en horas académicas.', 'Obligatorio'),
    ('Costo',             'Costo total del curso para la empresa.', 'Opcional'),
    ('Descripción',       'Contenido y objetivos del curso.', 'Opcional'),
    ('Estado',            'Planificado / En curso / Finalizado / Cancelado.', 'Automático'),
]

tbl = doc.add_table(rows=len(campos_curso)+1, cols=3)
set_col_widths(tbl, [4, 8, 3])
header_table_row(tbl, 0, ['Campo', 'Descripción', 'Requerido'])
for i, row in enumerate(campos_curso):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()
add_body(doc, 'Para inscribir empleados: desde el curso, haga clic en "Inscribir Empleado" y seleccione los participantes. Al finalizar el curso puede registrar si el empleado aprobó y su nota o certificado.')

# 3.7 Clima organizacional
add_h2(doc, '3.7 Clima Organizacional')

add_body(doc, 'Ruta: Talento Humano → Clima Organizacional')
add_body(doc,
    'Herramienta para medir la satisfacción y el ambiente laboral mediante encuestas '
    'periódicas. Los resultados se grafican para identificar áreas de mejora.')

add_image_placeholder(doc, 'Pantalla: Encuestas de Clima — resultados y gráficos de tendencia')

add_bullet(doc, 'Cree una encuesta con preguntas de escala (1-5 o 1-10) y preguntas abiertas.')
add_bullet(doc, 'Defina el período de respuesta (por ejemplo, 2 semanas).')
add_bullet(doc, 'Los empleados responden desde su perfil de forma anónima o identificada.')
add_bullet(doc, 'El sistema genera reportes con promedios por categoría y comparativo histórico.')

add_nota(doc,
    'Se recomienda realizar encuestas de clima al menos dos veces al año para monitorear '
    'tendencias y evaluar el impacto de las acciones de mejora implementadas.', AMARILLO)

# 3.8 Finiquitos
add_h2(doc, '3.8 Finiquitos')

add_body(doc, 'Ruta: Talento Humano → Finiquitos')
add_body(doc,
    'Calcula automáticamente la liquidación final del empleado al término de la relación '
    'laboral, cumpliendo con el Código del Trabajo ecuatoriano.')

add_image_placeholder(doc, 'Pantalla: Finiquito — cálculo y resumen de rubros')

add_h3(doc, 'Rubros calculados automáticamente')

rubros_fin = [
    ('Sueldo pendiente',          'Días trabajados del último mes parcial × (sueldo ÷ 30).'),
    ('Decimotercero pendiente',   'Proporcional al período no pagado (sueldo total ÷ 12 × meses/12).'),
    ('Decimocuarto pendiente',    'Proporcional al SBU vigente × meses trabajados en el año de referencia.'),
    ('Vacaciones pendientes',     'Días de vacaciones acumulados y no gozados × (sueldo ÷ 30).'),
    ('Fondo de reserva',          'Proporcional si no fue pagado mensualmente.'),
    ('Desahucio (art. 184)',       'Si aplica: 25% de la última remuneración mensual por año de servicio.'),
    ('Indemnización (art. 188)',   'Si aplica por despido intempestivo: según tabla legal del Código del Trabajo.'),
    ('Descuentos',                'Anticipos no descontados, préstamos pendientes y otros compromisos.'),
    ('Total a pagar',             'Suma de haberes menos descuentos.'),
]

tbl = doc.add_table(rows=len(rubros_fin)+1, cols=2)
set_col_widths(tbl, [5, 10.5])
header_table_row(tbl, 0, ['Rubro', 'Fórmula / Descripción'])
for i, row in enumerate(rubros_fin):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()
add_body(doc, 'Al guardar el finiquito, el empleado cambia automáticamente a estado "Inactivo" y se registra la fecha de salida.')
add_nota(doc,
    'El finiquito generado por el sistema es una herramienta de cálculo de referencia. '
    'El acta de finiquito oficial debe ser firmada ante el Inspector de Trabajo del MRL.',
    AMARILLO)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  4. MÓDULO DE NÓMINAS
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, '4. Módulo de Liquidación de Nóminas')

add_body(doc,
    'Este módulo maneja el ciclo completo de la nómina mensual: desde la creación del período '
    'hasta el cierre definitivo, incluyendo el cálculo del rol de pagos, anticipos, novedades, '
    'décimos, vacaciones y la contabilización automática.')

# 4.1 Períodos
add_h2(doc, '4.1 Períodos de Nómina')

add_body(doc, 'Ruta: Liquidación de Nóminas → Períodos')
add_body(doc,
    'Cada período representa un mes de nómina. El flujo de estados es '
    'Borrador → Liquidado. Solo puede existir un período en estado "Borrador" a la vez.')

add_image_placeholder(doc, 'Pantalla: Períodos de Nómina — listado con estado y acciones')

estados_periodo = [
    ('Borrador',   'El período está abierto. Se pueden generar y modificar roles, anticipos y novedades.'),
    ('Liquidado',  'Período cerrado definitivamente. Los roles fueron procesados, la contabilidad fue registrada. No se puede modificar.'),
]

tbl = doc.add_table(rows=len(estados_periodo)+1, cols=2)
set_col_widths(tbl, [3, 12.5])
header_table_row(tbl, 0, ['Estado', 'Descripción'])
for i, row in enumerate(estados_periodo):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()

add_h3(doc, 'Crear un Período')

add_bullet(doc, 'Haga clic en "Nuevo Período".')
add_bullet(doc, 'Ingrese el nombre (Ej.: Nómina Junio 2026), la fecha de inicio y la fecha de fin.')
add_bullet(doc, 'El sistema verifica que no exista otro período con las mismas fechas.')
add_bullet(doc, 'Guarde — el período queda en estado "Borrador".')

add_h3(doc, 'Liquidar un Período')

add_body(doc,
    'Al hacer clic en "Liquidar Período", el sistema ejecuta:')
add_bullet(doc, 'Verifica que todos los roles del período estén en estado "Aprobado".')
add_bullet(doc, 'Genera los asientos contables en LedgerPro (rol mensual + provisión de leyes sociales).')
add_bullet(doc, 'Cambia el estado del período a "Liquidado".')
add_bullet(doc, 'Bloquea cualquier modificación posterior al período.')

add_nota(doc,
    'Una vez liquidado, el período NO puede reabrirse. Verifique los valores del rol antes de liquidar.',
    AMARILLO, 'ADVERTENCIA')

# 4.2 Rol de pagos
add_h2(doc, '4.2 Rol de Pagos')

add_body(doc, 'Ruta: Liquidación de Nóminas → Períodos → (seleccione período) → Ver Rol')
add_body(doc,
    'El rol de pagos es el documento que detalla todos los ingresos y descuentos de cada '
    'empleado en el período. QuickInvoice lo calcula automáticamente.')

add_image_placeholder(doc, 'Pantalla: Rol de Pagos — cabecera del período y listado de empleados')

add_h3(doc, 'Generar el Rol')

add_body(doc, 'Desde el período en estado "Borrador":')
add_bullet(doc, 'Haga clic en "Generar Rol" — el sistema crea una línea por cada empleado activo.')
add_bullet(doc, 'Los conceptos configurados como "Aplica siempre" se incluyen automáticamente.')
add_bullet(doc, 'Las novedades activas del período se incluyen como descuentos automáticos.')
add_bullet(doc, 'Los anticipos ya liquidados en el período se descuentan automáticamente.')
add_bullet(doc, 'Se calculan: IESS personal, aporte patronal, y el neto a pagar.')

add_image_placeholder(doc, 'Pantalla: Detalle del Rol — ingresos, descuentos, IESS y neto por empleado')

add_h3(doc, 'Estructura del Rol por Empleado')

rol_estructura = [
    ('Sueldo base',           'ingreso',   'Remuneración mensual pactada.'),
    ('Horas extras 50%',      'ingreso',   'Monto calculado automáticamente desde las horas ingresadas.'),
    ('Horas extras 100%',     'ingreso',   'Monto calculado automáticamente desde las horas ingresadas.'),
    ('Horas nocturnas 25%',   'ingreso',   'Recargo nocturno calculado automáticamente.'),
    ('Otros ingresos',        'ingreso',   'Comisiones, bonos u otros rubros adicionales.'),
    ('IESS Personal (9.45%)', 'descuento', 'Calculado sobre el total de ingresos que afectan IESS.'),
    ('Anticipo quincena',     'descuento', 'Anticipo ya desembolsado en el período.'),
    ('Novedades / Préstamos', 'descuento', 'Cuotas de préstamos y descuentos recurrentes.'),
    ('Retención Renta',       'descuento', 'Retención en la fuente si el ingreso supera la base imponible.'),
    ('TOTAL INGRESOS',        '─',         'Suma de todos los conceptos de ingreso.'),
    ('TOTAL DESCUENTOS',      '─',         'Suma de todos los conceptos de descuento.'),
    ('NETO A PAGAR',          '─',         'Total Ingresos − Total Descuentos.'),
]

tbl = doc.add_table(rows=len(rol_estructura)+1, cols=3)
set_col_widths(tbl, [5, 2.5, 8])
header_table_row(tbl, 0, ['Concepto', 'Tipo', 'Descripción'])
for i, row in enumerate(rol_estructura):
    fill = AZUL_CLARO if row[0].startswith('TOTAL') or row[0] == 'NETO A PAGAR' else None
    data_row(tbl, i+1, list(row), fill)
add_table_border(tbl)

doc.add_paragraph()

add_h3(doc, 'Ingresar Horas Extras / Nocturnas')

add_body(doc,
    'Para los conceptos de horas extra y nocturnas, el modal de edición del rol '
    'solicita el número de HORAS trabajadas (no el monto). El sistema calcula automáticamente:')

add_image_placeholder(doc, 'Pantalla: Modal edición empleado — campo "Horas" para HRS_EXTRA_50')

add_bullet(doc, 'HRS_EXTRA_50:  (Sueldo ÷ 240) × 1.5 × horas ingresadas')
add_bullet(doc, 'HRS_EXTRA_100: (Sueldo ÷ 240) × 2.0 × horas ingresadas')
add_bullet(doc, 'HRS_NOCT_25:   (Sueldo ÷ 240) × 0.25 × horas ingresadas')
add_body(doc, 'El monto calculado aparece en modo lectura junto al input de horas.')

add_h3(doc, 'Modificar el Rol Manualmente')

add_body(doc,
    'Si necesita ajustar valores antes de aprobar el rol:')
add_bullet(doc, 'Haga clic en el ícono de edición del empleado.')
add_bullet(doc, 'Puede agregar conceptos adicionales, cambiar montos o eliminar líneas.')
add_bullet(doc, 'El sistema recalcula automáticamente el IESS personal y el neto.')
add_bullet(doc, 'Para conceptos de horas: ingrese siempre el número de horas, no el monto.')

add_h3(doc, 'Aprobar y Cerrar el Rol')

add_bullet(doc, 'Revise los valores en el resumen del período (totales por columna).')
add_bullet(doc, 'Si todo está correcto, apruebe el rol individual de cada empleado o use "Aprobar todo".')
add_bullet(doc, 'Una vez aprobados todos los roles, puede proceder a "Liquidar Período".')

add_nota(doc,
    'El sistema calcula el IESS patronal (11.15%) de manera separada — este monto NO aparece '
    'en el rol del empleado pero sí se registra en la contabilidad como carga patronal.',
    AMARILLO)

# 4.3 Anticipos
add_h2(doc, '4.3 Anticipos de Quincena')

add_body(doc, 'Ruta: Liquidación de Nóminas → Períodos → (período) → Anticipo')
add_body(doc,
    'Gestión de anticipos de quincena para empleados que solicitan un adelanto de su sueldo '
    'mensual. El sistema calcula el monto sugerido y descuenta automáticamente del rol.')

add_image_placeholder(doc, 'Pantalla: Anticipo de Quincena — listado de empleados con monto sugerido')

add_h3(doc, 'Proceso de Anticipo')

pasos_anticipo = [
    'Ingrese al período activo y haga clic en "Anticipo de Quincena".',
    'Haga clic en "Generar Líneas" — el sistema sugiere el monto para cada empleado según su configuración.',
    'Revise y ajuste los montos si es necesario. Puede dejar en cero los empleados que no tomarán anticipo.',
    'Haga clic en "Liquidar Anticipo" para procesar el pago y registrar el asiento contable.',
    'El anticipo liquidado se descuenta automáticamente del rol mensual del período actual.',
]

for i, paso in enumerate(pasos_anticipo):
    add_bullet(doc, f'Paso {i+1}: {paso}')

add_h3(doc, 'Configuración del Monto de Anticipo por Empleado')

add_body(doc, 'En la ficha del empleado → Pestaña Anticipo:')

anticipo_tipos = [
    ('Porcentaje',    'El anticipo es un % del sueldo base (Ej.: 40% → empleado con $800 recibe $320).'),
    ('Monto fijo',    'El anticipo es siempre el mismo valor independientemente del sueldo.'),
    ('Sin anticipo',  'El empleado no tiene anticipo configurado (monto = 0 en la generación automática).'),
]

tbl = doc.add_table(rows=len(anticipo_tipos)+1, cols=2)
set_col_widths(tbl, [4, 11.5])
header_table_row(tbl, 0, ['Tipo', 'Comportamiento'])
for i, row in enumerate(anticipo_tipos):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()
add_nota(doc,
    'Si el anticipo ya fue liquidado y necesita deshacerlo (antes de liquidar el período), '
    'use el botón "Deshacer Liquidación". Esto revierte el asiento contable y permite re-editar.', AMARILLO)

# 4.4 Novedades
add_h2(doc, '4.4 Novedades y Descuentos Recurrentes')

add_body(doc, 'Ruta: Liquidación de Nóminas → Novedades / Descuentos')
add_body(doc,
    'Las novedades son descuentos que se repiten en múltiples períodos sin necesidad de '
    'ingresarlos manualmente cada mes. Son ideales para préstamos y descuentos fijos.')

add_image_placeholder(doc, 'Pantalla: Novedades — lista con tipo, monto, saldo y barra de progreso')

tipos_novedad = [
    ('Descuento Variable',   'El usuario ingresa el monto cada mes al revisar el rol. Útil para descuentos que cambian.',
                             'No aplica', 'Monto = 0, usuario lo edita'),
    ('Descuento Fijo',       'Se descuenta siempre el mismo monto mensual hasta desactivarse manualmente.',
                             'Monto fijo', 'Cuota fija mensual'),
    ('Préstamo por Cuota',   'Préstamo con cuota mensual fija. Se descuenta hasta agotar el saldo.',
                             'Saldo inicial + cuota mensual', 'Cuota hasta saldo = 0'),
    ('Préstamo a Plazos',    'Préstamo dividido en N meses iguales. La cuota = saldo inicial ÷ N meses.',
                             'Saldo inicial + N meses', 'Cuota = saldo ÷ N'),
]

tbl = doc.add_table(rows=len(tipos_novedad)+1, cols=4)
set_col_widths(tbl, [3.5, 5.5, 3.5, 3])
header_table_row(tbl, 0, ['Tipo', 'Cuándo usar', 'Datos requeridos', 'Comportamiento'])
for i, row in enumerate(tipos_novedad):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()
add_image_placeholder(doc, 'Pantalla: Modal Nueva Novedad — tipo de novedad, monto y fecha inicio')

add_h3(doc, 'Saldos de Préstamos')

add_body(doc,
    'Al cerrar cada período, el sistema actualiza automáticamente el saldo pendiente de cada préstamo:')
add_bullet(doc, 'saldo_pendiente = saldo_pendiente − cuota_del_período')
add_bullet(doc, 'n_cuotas_pagadas = n_cuotas_pagadas + 1')
add_bullet(doc, 'Si saldo_pendiente ≤ 0: la novedad se marca automáticamente como inactiva.')

add_body(doc,
    'La barra de progreso visible en el listado muestra el porcentaje del préstamo ya pagado:')
add_bullet(doc, 'Verde completo = pagado 100%.')
add_bullet(doc, 'Naranja parcial = en proceso.')

# 4.5 Décimos
add_h2(doc, '4.5 Liquidación de Décimos (13.º y 14.º Sueldo)')

add_body(doc, 'Ruta: Liquidación de Nóminas → Liquidación Décimos')
add_body(doc,
    'Calcula y registra el pago del decimotercero (proporcional en diciembre) y del '
    'decimocuarto (proporcional, pagadero según la región del empleado).')

add_image_placeholder(doc, 'Pantalla: Liquidación de Décimos — tabla por empleado con rubros y totales')

add_h3(doc, 'Decimotercero (13.º sueldo)')

add_body(doc, 'Fórmula según el Código del Trabajo ecuatoriano:')
add_bullet(doc, 'Período: 1 de diciembre del año anterior al 30 de noviembre del año en curso.')
add_bullet(doc, 'Base de cálculo: suma de todos los ingresos percibidos en el período.')
add_bullet(doc, 'Fórmula: Total ingresos del período ÷ 12.')
add_bullet(doc, 'Puede pagarse de forma acumulada o mensualizada (según elección del empleado en el MRL).')

add_h3(doc, 'Decimocuarto (14.º sueldo)')

add_body(doc, 'Fórmula:')
add_bullet(doc, 'Período Sierra/Amazonía: 1 de agosto año anterior al 31 de julio año en curso.')
add_bullet(doc, 'Período Costa/Galápagos: 1 de marzo año anterior al 28 de febrero año en curso.')
add_bullet(doc, 'Fórmula: SBU vigente × meses trabajados en el período ÷ 12.')
add_bullet(doc, 'Para empleados de tiempo parcial: proporcional a las horas trabajadas.')

add_nota(doc,
    'Si el empleado eligió mensualizar el 13.º o 14.º (declaró en el MRL), el pago se hace '
    'mensualmente con el rol. En ese caso registre el pago desde el Rol como concepto adicional, '
    'no desde este módulo de liquidación.', AMARILLO)

# 4.6 Vacaciones
add_h2(doc, '4.6 Liquidación de Vacaciones')

add_body(doc, 'Ruta: Liquidación de Nóminas → Liquidación Vacaciones')
add_body(doc,
    'Calcula los días de vacaciones acumulados y el valor a pagar cuando el empleado '
    'toma vacaciones o cuando se liquidan al finalizar la relación laboral.')

add_image_placeholder(doc, 'Pantalla: Vacaciones — tabla con días acumulados y valor a pagar por empleado')

add_bullet(doc, 'Derecho: 15 días de vacaciones por cada año completo de trabajo (Art. 69 CT).')
add_bullet(doc, 'A partir del quinto año: un día adicional por cada año extra, hasta un máximo de 15 días adicionales.')
add_bullet(doc, 'Fórmula: (Sueldo ÷ 30) × días de vacaciones.')
add_bullet(doc, 'Los días acumulados se calculan desde la fecha de ingreso hasta la fecha de corte.')
add_bullet(doc, 'Puede registrar vacaciones gozadas (sin pago) o vacaciones pagadas.')

# 4.7 Estado económico
add_h2(doc, '4.7 Estado Económico / Capacidad de Pago')

add_body(doc, 'Ruta: Liquidación de Nóminas → Estado Económico')
add_body(doc,
    'Panel de análisis financiero que muestra la carga laboral total de la empresa '
    'y permite evaluar la capacidad de pago de la nómina mes a mes.')

add_image_placeholder(doc, 'Pantalla: Estado Económico — gráficos de carga laboral y tendencia')

add_bullet(doc, 'Carga mensual total: neto a pagar + IESS patronal + provisiones del mes.')
add_bullet(doc, 'Comparativo entre períodos: evolución de la nómina en el tiempo.')
add_bullet(doc, 'Desglose por sección/departamento: qué área representa mayor costo laboral.')
add_bullet(doc, 'Indicadores de alertas: empleados con anticipos que superan el 50% del sueldo.')

# 4.8 Reportes
add_h2(doc, '4.8 Reportes de Nómina')

add_body(doc, 'Ruta: Liquidación de Nóminas → Períodos → (período) → Reportes')
add_body(doc,
    'Genera reportes detallados del rol para el período seleccionado, listos para imprimir '
    'o exportar como respaldo y para presentación ante entidades de control.')

add_image_placeholder(doc, 'Pantalla: Reporte de Nómina — vista previa para impresión')

reportes_disp = [
    ('Rol consolidado',       'Tabla con todos los empleados, ingresos, descuentos y neto. Firma del empleado.'),
    ('Reporte IESS',          'Detalle del IESS personal y patronal por empleado para el formulario de la planilla.'),
    ('Reporte anticipos',     'Detalle de anticipos del período con firmas de recibido.'),
    ('Rol individual',        'Comprobante de pago individual por empleado (rol de pagos).'),
    ('Décimos acumulados',    'Estado actual de la provisión de décimos por empleado.'),
]

tbl = doc.add_table(rows=len(reportes_disp)+1, cols=2)
set_col_widths(tbl, [5, 10.5])
header_table_row(tbl, 0, ['Reporte', 'Contenido'])
for i, row in enumerate(reportes_disp):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  5. CONTABILIDAD AUTOMÁTICA
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, '5. Contabilidad Automática de Nómina')

add_body(doc,
    'QuickInvoice genera automáticamente los asientos contables en LedgerPro al liquidar '
    'cada período o al procesar anticipos, sin necesidad de intervención manual.')

add_nota(doc,
    'Requisito previo: las cuentas contables deben estar configuradas en "Cuentas Contables '
    'Nómina" (sección 2.3). Si alguna cuenta falta, la contabilización falla y se muestra un error.', AMARILLO)

add_h2(doc, '5.1 Asiento del Rol Mensual')

add_body(doc,
    'Al liquidar el período, se generan DOS asientos contables separados en LedgerPro:')

add_h3(doc, 'Asiento 1 — Rol de Pagos')

add_body(doc, 'Referencia: rol_mensual_{periodoId}')

asiento1 = [
    ('Gasto Sueldos',             'DEBE',  'Total remuneraciones brutas del período.'),
    ('Gasto IESS Patronal',       'DEBE',  'Aporte patronal 11.15% sobre base IESS.'),
    ('Sueldos por Pagar',         'HABER', 'Neto total a pagar a empleados.'),
    ('IESS por Pagar',            'HABER', 'IESS personal (9.45%) + IESS patronal (11.15%).'),
    ('Anticipo Empleados',        'HABER', 'Total anticipos ya desembolsados a descontar.'),
]

tbl = doc.add_table(rows=len(asiento1)+1, cols=3)
set_col_widths(tbl, [6, 2, 7.5])
header_table_row(tbl, 0, ['Cuenta', 'DEBE / HABER', 'Valor'])
for i, row in enumerate(asiento1):
    fill = AZUL_CLARO if row[1] == 'DEBE' else VERDE_CLARO
    data_row(tbl, i+1, list(row), fill)
add_table_border(tbl)

doc.add_paragraph()

add_h3(doc, 'Asiento 2 — Provisión de Leyes Sociales')

add_body(doc, 'Referencia: provision_leyes_{periodoId}')

asiento2 = [
    ('Gasto Décimo Tercero',      'DEBE',  'Provisión mensual = base IESS ÷ 12.'),
    ('Gasto Décimo Cuarto',       'DEBE',  'Provisión mensual = SBU ÷ 12.'),
    ('Gasto Vacaciones',          'DEBE',  'Provisión mensual = base IESS × 4.17%.'),
    ('Gasto Fondo de Reserva',    'DEBE',  'Provisión mensual = base IESS × 8.33% (a partir del 2.º año).'),
    ('Prov. Décimo Tercero',      'HABER', 'Pasivo acumulado por pagar.'),
    ('Prov. Décimo Cuarto',       'HABER', 'Pasivo acumulado por pagar.'),
    ('Prov. Vacaciones',          'HABER', 'Pasivo acumulado por pagar.'),
    ('Prov. Fondo de Reserva',    'HABER', 'Pasivo acumulado por pagar.'),
]

tbl = doc.add_table(rows=len(asiento2)+1, cols=3)
set_col_widths(tbl, [6, 2, 7.5])
header_table_row(tbl, 0, ['Cuenta', 'DEBE / HABER', 'Valor'])
for i, row in enumerate(asiento2):
    fill = AZUL_CLARO if row[1] == 'DEBE' else VERDE_CLARO
    data_row(tbl, i+1, list(row), fill)
add_table_border(tbl)

doc.add_paragraph()
add_nota(doc,
    'La provisión descuenta lo ya pagado en décimos, vacaciones y fondo de reserva '
    'durante el período para evitar duplicar el gasto. Solo se provisiona la diferencia pendiente.',
    VERDE_CLARO)

add_h2(doc, '5.2 Asiento del Anticipo')

add_body(doc, 'Al liquidar un anticipo de quincena, se genera un asiento contable:')

asiento_ant = [
    ('Anticipo Empleados',    'DEBE',  'Total anticipos desembolsados en el período.'),
    ('Bancos / Caja',         'HABER', 'Salida de efectivo o transferencia bancaria.'),
]

tbl = doc.add_table(rows=len(asiento_ant)+1, cols=3)
set_col_widths(tbl, [6, 2, 7.5])
header_table_row(tbl, 0, ['Cuenta', 'DEBE / HABER', 'Valor'])
for i, row in enumerate(asiento_ant):
    fill = AZUL_CLARO if row[1] == 'DEBE' else VERDE_CLARO
    data_row(tbl, i+1, list(row), fill)
add_table_border(tbl)

add_h2(doc, '5.3 Anulación de Asientos')

add_body(doc,
    'Si se deshace la liquidación de un período o anticipo, el sistema anula '
    'automáticamente TODOS los asientos contables generados:')
add_bullet(doc, 'Deshacen liquidación del período → anula asiento rol mensual + asiento provisión leyes.')
add_bullet(doc, 'Deshacer anticipo → anula asiento del anticipo.')
add_bullet(doc, 'Los asientos anulados quedan registrados en LedgerPro con estado "Anulado" para auditoría.')

add_image_placeholder(doc, 'Pantalla LedgerPro: Asientos del período — rol mensual y provisión')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  6. PERMISOS
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, '6. Permisos de Usuario')

add_body(doc,
    'El acceso a las diferentes secciones del módulo está controlado por permisos '
    'individuales por usuario, configurables desde Configuración → Permisos de Usuario.')

add_image_placeholder(doc, 'Pantalla: Permisos de Usuario — módulo Talento Humano y Nóminas')

permisos = [
    ('perm_th_empleados',          'Gestión de empleados, vacantes, evaluaciones, capacitación, clima, dashboard, finiquitos, checklists.'),
    ('perm_th_estructura',         'Estructura organizativa (secciones y cargos).'),
    ('perm_th_rol_nomina',         'Períodos, rol de pagos, reportes, anticipos, novedades, décimos, vacaciones, capacidad de pago.'),
    ('perm_th_conceptos_nomina',   'Configuración de conceptos de nómina (ingresos/descuentos).'),
    ('perm_th_nomina_parametros',  'Parámetros de nómina (SBU, IESS%) y cuentas contables de nómina.'),
]

tbl = doc.add_table(rows=len(permisos)+1, cols=2)
set_col_widths(tbl, [5.5, 10])
header_table_row(tbl, 0, ['Permiso', 'Acceso que otorga'])
for i, row in enumerate(permisos):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()

add_h2(doc, 'Roles sugeridos')

roles_sug = [
    ('Administrador Nómina',    'perm_th_empleados + perm_th_rol_nomina + perm_th_estructura',              'Puede operar todo el módulo excepto cambiar configuración base.'),
    ('Analista de Nómina',      'perm_th_rol_nomina',                                                       'Solo puede trabajar con períodos, roles, anticipos y liquidaciones.'),
    ('Jefe de RRHH',            'perm_th_empleados + perm_th_estructura',                                   'Gestión de personal, evaluaciones y onboarding. Sin acceso a nómina.'),
    ('Gerente General',         'Todos los permisos',                                                        'Acceso completo a todos los módulos.'),
    ('Contador',                'perm_th_rol_nomina + perm_th_nomina_parametros + perm_th_conceptos_nomina','Configura parámetros y supervisa contabilidad de nómina.'),
]

tbl = doc.add_table(rows=len(roles_sug)+1, cols=3)
set_col_widths(tbl, [4, 5.5, 6])
header_table_row(tbl, 0, ['Rol sugerido', 'Permisos recomendados', 'Acceso'])
for i, row in enumerate(roles_sug):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  7. FLUJO MENSUAL
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, '7. Flujo de Trabajo Mensual Recomendado')

add_body(doc,
    'Siga este flujo cada mes para garantizar una nómina correcta, oportuna y '
    'correctamente contabilizada:')

flujo = [
    ('1', 'Semana 1 — Quincenal',      'Anticipo',
     'Crear/liquidar el anticipo de quincena. Verifica que los montos sean correctos antes de liquidar.'),
    ('2', 'Semana 2',                   'Novedades',
     'Revisar y actualizar novedades: agregar nuevos préstamos, desactivar novedades saldadas.'),
    ('3', 'Semana 2-3',                 'Generar Rol',
     'Desde el período activo, hacer clic en "Generar Rol". El sistema crea las líneas de todos los empleados.'),
    ('4', 'Semana 3',                   'Revisar Rol',
     'Ingresar las horas extras/nocturnas de cada empleado. Verificar ingresos variables y descuentos especiales.'),
    ('5', 'Semana 3',                   'Aprobar Rol',
     'Revisar los totales del período. Aprobar el rol cuando todo esté correcto.'),
    ('6', 'Último día del mes',         'Liquidar Período',
     'Hacer clic en "Liquidar Período". El sistema genera los asientos contables en LedgerPro automáticamente.'),
    ('7', 'Primer día del mes siguiente', 'Nuevo Período',
     'Crear el nuevo período de nómina para el mes siguiente.'),
    ('8', 'Según calendario',           'Décimos / Vacaciones',
     'Procesar la liquidación de décimos y/o vacaciones cuando corresponda (diciembre para 13.º, fecha de corte para 14.º).'),
]

tbl = doc.add_table(rows=len(flujo)+1, cols=4)
set_col_widths(tbl, [0.8, 3.5, 3, 8.2])
header_table_row(tbl, 0, ['#', 'Cuándo', 'Acción', 'Detalle'])
for i, row in enumerate(flujo):
    data_row(tbl, i+1, list(row))
add_table_border(tbl)

doc.add_paragraph()
add_image_placeholder(doc, 'Diagrama de flujo mensual de nómina (opcional: insertar flujograma)')

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  8. PREGUNTAS FRECUENTES
# ════════════════════════════════════════════════════════════════════════════

add_h1(doc, '8. Preguntas Frecuentes')

faqs = [
    ('¿Por qué el sistema no me deja generar el rol?',
     'Verifique que: (a) exista al menos un empleado activo, (b) el período esté en estado "Borrador", '
     '(c) las fechas del período sean válidas.'),

    ('¿El IESS patronal aparece en el recibo del empleado?',
     'No. El IESS patronal (11.15%) es una carga del empleador y NO se refleja en el rol del empleado. '
     'Sí se registra en el asiento contable como gasto patronal.'),

    ('¿Cómo corrijo un error en el rol después de liquidar el período?',
     'Un período liquidado NO puede modificarse directamente. Contacte al administrador para '
     'gestionar el ajuste en el período siguiente como diferencia de nómina, o solicite la anulación '
     'del período (requiere permisos especiales de administrador).'),

    ('¿El sistema calcula retención en la fuente (impuesto a la renta)?',
     'El sistema incluye el concepto RETENCION_RF como descuento. Sin embargo, el monto de retención '
     'debe calcularse según la tabla del SRI vigente para el ejercicio fiscal. Puede ingresarlo manualmente '
     'en el rol de cada empleado que supere la base imponible anual.'),

    ('¿Los anticipos de quincena afectan el IESS?',
     'No. Los anticipos son simplemente adelantos del sueldo ya ganado. No generan ningún tipo de '
     'aporte adicional al IESS.'),

    ('¿Cómo agrego un nuevo empleado a mitad de mes?',
     'Cree la ficha del empleado con la fecha de ingreso correcta. Al generar el rol, el sistema '
     'incluirá al empleado. Debe ajustar manualmente el sueldo proporcional (días trabajados × sueldo ÷ 30) '
     'editando la línea del empleado en el rol.'),

    ('¿Qué pasa si el candidato seleccionado ya es empleado?',
     'Al cambiar el estado de un candidato a "Contratado", puede vincularlo directamente con un '
     'registro de empleado existente o crear uno nuevo. El historial del proceso de reclutamiento '
     'queda asociado al empleado.'),

    ('¿Cómo elimino una novedad (préstamo) que ya fue saldado?',
     'El sistema desactiva automáticamente los préstamos cuando el saldo llega a cero. '
     'Si necesita desactivarlo manualmente antes, use el botón "Desactivar" en el listado de novedades. '
     'Las novedades desactivadas no se incluyen en nuevos roles.'),

    ('¿Puedo tener dos períodos de nómina activos al mismo tiempo?',
     'No. El sistema solo permite un período en estado "Borrador" a la vez. Debe liquidar el período '
     'actual antes de crear uno nuevo.'),

    ('¿La evaluación 360° es anónima?',
     'Sí. Los evaluadores en la modalidad 360° no son identificables por el evaluado. '
     'Solo el administrador del sistema puede ver quién respondió cada formulario.'),

    ('¿Qué pasa con los asientos contables si deshago una liquidación?',
     'El sistema anula automáticamente todos los asientos generados (rol mensual + provisión de leyes). '
     'Los asientos anulados quedan en LedgerPro con estado "Anulado" para auditoría.'),

    ('¿Cómo actualizo el SBU cada año?',
     'Vaya a Liquidación de Nóminas → Parámetros de Nómina, edite el valor del SBU y guarde. '
     'El cambio aplica inmediatamente a todos los roles nuevos que se generen desde ese momento.'),
]

for i, (preg, resp) in enumerate(faqs):
    p = doc.add_paragraph()
    r = p.add_run(f'P{i+1}: {preg}')
    r.font.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    p.paragraph_format.space_before = Pt(8)

    p2 = doc.add_paragraph()
    r2 = p2.add_run(resp)
    r2.font.size = Pt(10.5)
    p2.paragraph_format.left_indent = Cm(0.5)
    p2.paragraph_format.space_after = Pt(4)

page_break(doc)

# ════════════════════════════════════════════════════════════════════════════
#  CONTRAPORTADA
# ════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Billennium Systems')
r.font.size  = Pt(16)
r.font.bold  = True
r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('QuickInvoice — Sistema de Gestión Empresarial')
r.font.size  = Pt(12)
r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run(f'Manual de Usuario — Talento Humano y Nóminas  |  v1.0  |  {datetime.date.today().strftime("%B %Y")}')
r.font.size  = Pt(10)
r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

add_image_placeholder(doc, 'Logo o imagen de pie de página (opcional)', '8 cm', '2 cm')

# ════════════════════════════════════════════════════════════════════════════
#  GUARDAR
# ════════════════════════════════════════════════════════════════════════════

output = r'C:\Billenniumweb2026\websitebillennium-main\App\Proyecto QuickInvoice\Manuales\Manual_TalentoHumano_Nominas_v1.docx'
doc.save(output)
print(f'Documento guardado en: {output}')
